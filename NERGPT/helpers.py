from transformers import AutoTokenizer, AutoModelForSequenceClassification

import torch
import docx
from PyPDF2 import PdfReader, PdfFileWriter

import openai

TOKENIZERS_PARALLELISM=False



tokenizer = AutoTokenizer.from_pretrained(
    "ivanlau/language-detection-fine-tuned-on-xlm-roberta-base")

model = AutoModelForSequenceClassification.from_pretrained(
    "ivanlau/language-detection-fine-tuned-on-xlm-roberta-base")

def remove_special_characters(text: str):
    text = text.replace('\t', " ")
    text = text.replace('/', "")
    # text = text.replace("§", " ")
    return text

def split_text(text: str) -> list[str]:
    """
    Handles line breaks, and cleans the text like that. 
    """
    sections = text.split("\n")

    def handle_line_breaks():
        final = []
        i = 0
        for i in range(len(sections)):
            if sections[i]:
                # print("current section ", sections[i])
                if len(final) > 0:
                    # print("last section ",final[-1])
                    # print("last word ",final[-1][-1] )
                    if final[-1][-1] == "-":  # The last accepted string ends with -
                        # Check if the second to last character is a letter
                        if final[-1][-2].isalpha():
                            last_word = final[-1][-1][:-1] + sections[i][0]
                            final[-1] = final[-1][:-1] + last_word + \
                                remove_special_characters(sections[i][1:])
                            continue
                to_add = remove_special_characters(sections[i])
                if to_add:
                    final.append(to_add.strip())
        return final

    return handle_line_breaks()

def get_german_sentences(sentences):
    germanSentences = []
    for sentence in sentences:

        if(len(sentence.split())>1):
            inputs = tokenizer(sentence, return_tensors="pt")
            with torch.no_grad():
                logits = model(**inputs).logits
                predicted_class_id = logits.argmax().item()

                logits = torch.softmax(logits, dim=1)
                predicted_class_value = logits[0][predicted_class_id].item()
                label = model.config.id2label[predicted_class_id]
            if label == "German" and predicted_class_value >= 0.95:
                germanSentences.append(sentence)
        else:
            germanSentences.append(sentence)
    return '. '.join(germanSentences)


def readDocument(filename):
    with open(filename) as file:
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        fullText = '\n'.join(fullText)
        return split_text(fullText)


def processText(filename):
    raw_text = readDocument(filename)
    return get_german_sentences(raw_text)

#1- Tried getting german sentences directly from the text. The problem is that Czech and German get micex, might be hard to bring them back together
def get_paragraphs(filename:str)->list[list[str]]:
    german_sentences = processText(filename)
    return german_sentences



def get_first_page(filepath:str):
    """
    Parses the first page of the document, for extraction
    """
    reader = PdfReader(filepath)
    page = reader.pages[0]
    text = page.extract_text()
    return text

        